from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Dict, Iterable, Set

from docx import Document


SPANISH_MONTHS = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


def format_spanish_long(d: date) -> str:
    return f"{d.day} de {SPANISH_MONTHS[d.month]} de {d.year}"


_filename_bad = re.compile(r'[<>:"/\\|?*\x00-\x1F]+')


def safe_filename(name: str) -> str:
    s = name.strip()
    s = _filename_bad.sub("_", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:180] if len(s) > 180 else s


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    text = paragraph.text
    new_text = text
    for k, v in mapping.items():
        new_text = new_text.replace(k, v)

    if new_text == text:
        return

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def replace_placeholders(doc: Document, mapping: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def _collect_all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)

    return "\n".join(parts)


_placeholder_re = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def find_placeholders_in_template(template_bytes: bytes) -> Set[str]:
    doc = Document(BytesIO(template_bytes))
    text = _collect_all_text(doc)
    return set(_placeholder_re.findall(text))


def assert_required_placeholders(
    template_bytes: bytes,
    required: Iterable[str],
) -> None:
    found = find_placeholders_in_template(template_bytes)
    missing = [r for r in required if r not in found]
    if missing:
        raise RuntimeError(
            "Template missing required placeholders: "
            + ", ".join(f"{{{{{m}}}}}" for m in missing)
        )


@dataclass(frozen=True)
class DocContext:
    today: date
    code: str
    worker_name_upper: str
    incident_date: date
    observations: str


def render_docx(template_bytes: bytes, ctx: DocContext) -> bytes:
    doc = Document(BytesIO(template_bytes))

    mapping = {
        "{{today}}": format_spanish_long(ctx.today),
        "{{code}}": ctx.code,
        "{{name}}": ctx.worker_name_upper,
        "{{incident_date}}": format_spanish_long(ctx.incident_date),
        "{{observations}}": ctx.observations,
    }

    replace_placeholders(doc, mapping)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_output_filename(
    *,
    company_client_name: str,
    code: str,
    worker_full_name: str,
    worker_national_id: str,
    incident_type_code: str,
) -> str:
    base = (
        f"{company_client_name}__{incident_type_code}__{code}__"
        f"{worker_full_name}__{worker_national_id}.docx"
    )
    return safe_filename(base)


def save_bytes(folder: str, filename: str, content: bytes) -> str:
    out_dir = Path(folder)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / filename
    out_path.write_bytes(content)
    return str(out_path)